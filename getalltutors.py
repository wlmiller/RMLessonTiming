import os
from parseOSfile import parseOSfile
import re
from docx import *

contentdir = "C://Users/nmiller.RMCITY/Desktop/svn/"

for s in [f for f in os.listdir(contentdir) if re.match('0[0-5]-',f)]:
	for l in os.listdir(contentdir + '/' + s):
		lesson = l.split('-')[0]
		osfn = contentdir + '/' + s + '/' + l + '/3-OS/' +  lesson + '.docx'

		if os.path.exists(osfn):
			osfile = Document(osfn)
			tutor = "????"
			for par in osfile.paragraphs:
				if 'Martin' in par.text:
					tutor = 'Martin'
					break
				elif 'Stephanie' in par.text:
					tutor = 'Stephanie'
					break
				elif 'Angela' in par.text:
					tutor = 'Angela'
					break
			if tutor == '????':
				for tab in osfile.tables:
					for col in tab.columns:
						for cell in col.cells:
							for par in cell.paragraphs:
								if 'Martin' in par.text:
									tutor = 'Martin'
									break
								elif 'Stephanie' in par.text:
									tutor = 'Stephanie'
									break
								elif 'Angela' in par.text:
									tutor = 'Angela'
								break
				
			print "6-" + lesson + "\t" + tutor 
